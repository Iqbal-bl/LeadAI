"""
Audience building — turn an uploaded file (or a filter) into contactable rows.

THE HARD PART IS NOT PARSING, IT IS THE MESS
--------------------------------------------
Real customer lists arrive as: a CSV exported from Tally, an XLSX where row 1 is
a merged title and row 3 is the header, a DOCX someone pasted a table into, and
phone numbers written as `9876543210`, `+91 98765 43210`, `91-9876543210`,
`09876543210` and `9.87654321e+09` (Excel helpfully turned it into a float).
All five are the same person. This module normalises them, so a campaign does
not silently skip a third of the list or dial the same person twice.

WHAT IT GUARANTEES
------------------
* Header detection — finds the real header row rather than assuming row 1.
* Fuzzy column mapping — "Mobile No.", "contact number", "whatsapp" all map to
  phone; an explicit column_map always overrides the guess.
* E.164 normalisation with a configurable default country (India), via the
  `phonenumbers` library the outbound app already ships.
* De-duplication on the keyed phone fingerprint, counted and reported so the
  user sees "4,812 rows, 4,690 valid, 118 duplicates, 4 invalid" BEFORE sending.
* Every unmapped spreadsheet column is preserved in FieldsJson, so a message
  template can reference {{policy_number}} without a schema change.
"""
from __future__ import annotations

import csv
import io
import logging
import re
from dataclasses import dataclass, field
from typing import Any, Iterable

from ..config import settings
from ..security import encrypt_pii, mask_phone, phone_fingerprint

logger = logging.getLogger(__name__)

DEFAULT_REGION = "IN"

# Fuzzy header matching. Order matters: the first pattern that matches a header
# claims it, so `whatsapp` is checked before the looser `phone`.
FIELD_PATTERNS: list[tuple[str, tuple[str, ...]]] = [
    ("whatsapp", ("whatsapp", "wa number", "wa_no", "wanumber")),
    ("phone", ("phone", "mobile", "contact number", "contact no", "cell",
               "msisdn", "number", "tel", "mob")),
    ("email", ("email", "e-mail", "mail id", "emailid")),
    ("name", ("name", "customer", "client name", "full name", "contact name",
              "first name", "fname")),
    ("company", ("company", "organisation", "organization", "firm", "business")),
    ("city", ("city", "location", "town")),
    ("language", ("language", "lang", "preferred language")),
]


@dataclass
class ParsedRow:
    row_number: int
    name: str | None = None
    phone: str | None = None       # normalised E.164, or None
    email: str | None = None
    whatsapp: str | None = None
    country_code: str | None = None
    is_valid: bool = True
    invalid_reason: str | None = None
    fields: dict[str, Any] = field(default_factory=dict)


@dataclass
class ParseResult:
    rows: list[ParsedRow] = field(default_factory=list)
    headers: list[str] = field(default_factory=list)
    column_map: dict[str, str] = field(default_factory=dict)
    total: int = 0
    valid: int = 0
    invalid: int = 0
    duplicates: int = 0
    warnings: list[str] = field(default_factory=list)

    @property
    def sample(self) -> list[dict]:
        """First few rows, masked — what the UI shows on the preview screen."""
        return [
            {
                "row": r.row_number,
                "name": r.name,
                "phone": mask_phone(r.phone),
                "email": r.email,
                "valid": r.is_valid,
                "reason": r.invalid_reason,
            }
            for r in self.rows[:10]
        ]


# =========================================================================== #
# phone normalisation
# =========================================================================== #
def normalise_phone(raw: Any, region: str = DEFAULT_REGION) -> tuple[str | None, str | None]:
    """Return (e164, error). Handles the Excel-float and leading-zero cases."""
    if raw is None:
        return None, "missing"
    text = str(raw).strip()
    if not text or text.lower() in ("nan", "none", "null", "-"):
        return None, "missing"

    # Excel turned a 10-digit number into 9876543210.0 or 9.87654321E+09.
    if re.fullmatch(r"\d+\.\d+([eE][+-]?\d+)?", text):
        try:
            text = f"{int(float(text))}"
        except (ValueError, OverflowError):
            return None, "unreadable"

    text = re.sub(r"[^\d+]", "", text)
    if not text:
        return None, "no digits"

    try:
        import phonenumbers

        candidate = text if text.startswith("+") else text
        parsed = phonenumbers.parse(candidate, region)
        if not phonenumbers.is_valid_number(parsed):
            # A leading 0 is a national trunk prefix; strip and retry once.
            stripped = text.lstrip("0")
            if stripped != text:
                parsed = phonenumbers.parse(stripped, region)
            if not phonenumbers.is_valid_number(parsed):
                return None, "not a valid number"
        return phonenumbers.format_number(
            parsed, phonenumbers.PhoneNumberFormat.E164
        ), None
    except ImportError:  # pragma: no cover
        digits = re.sub(r"\D", "", text)
        if len(digits) < 10:
            return None, "too short"
        if not text.startswith("+"):
            digits = digits[-10:]
            return f"+91{digits}", None
        return f"+{digits}", None
    except Exception:  # noqa: BLE001
        return None, "unparseable"


def _valid_email(value: Any) -> str | None:
    text = str(value or "").strip()
    return text if re.fullmatch(r"[^@\s]+@[^@\s]+\.[A-Za-z]{2,}", text) else None


# =========================================================================== #
# header detection + mapping
# =========================================================================== #
def guess_column_map(headers: Iterable[str]) -> dict[str, str]:
    """{'phone': 'Mobile No.', 'name': 'Customer'} — best-effort, overridable."""
    mapping: dict[str, str] = {}
    claimed: set[str] = set()
    for canonical, patterns in FIELD_PATTERNS:
        for header in headers:
            if header in claimed:
                continue
            lowered = str(header or "").strip().lower()
            if not lowered:
                continue
            if any(p in lowered for p in patterns):
                mapping[canonical] = header
                claimed.add(header)
                break
    return mapping


def _find_header_row(matrix: list[list[Any]], limit: int = 10) -> int:
    """Pick the row that looks most like a header.

    Score = how many cells are short non-numeric strings that match a known
    field pattern. A merged title row ("CUSTOMER LIST 2026") scores 0; the real
    header row scores 2+.
    """
    best_index, best_score = 0, -1
    for index, row in enumerate(matrix[:limit]):
        cells = [str(c).strip().lower() for c in row if str(c or "").strip()]
        if not cells:
            continue
        score = sum(
            1
            for cell in cells
            for _, patterns in FIELD_PATTERNS
            if any(p in cell for p in patterns)
        )
        score += sum(1 for cell in cells if not cell.replace(".", "").isdigit()) * 0.1
        if score > best_score:
            best_index, best_score = index, score
    return best_index


# =========================================================================== #
# file readers
# =========================================================================== #
def _read_matrix(filename: str, content_type: str, blob: bytes) -> list[list[Any]]:
    """Every supported format collapses to a list of rows of cells."""
    name = (filename or "").lower()

    if name.endswith((".xlsx", ".xlsm", ".xltx")) or "spreadsheet" in (content_type or ""):
        from openpyxl import load_workbook

        workbook = load_workbook(io.BytesIO(blob), read_only=True, data_only=True)
        sheet = workbook.active
        return [list(row) for row in sheet.iter_rows(values_only=True)]

    if name.endswith(".xls"):
        import xlrd

        book = xlrd.open_workbook(file_contents=blob)
        sheet = book.sheet_by_index(0)
        return [sheet.row_values(r) for r in range(sheet.nrows)]

    if name.endswith(".docx") or "wordprocessing" in (content_type or ""):
        return _read_docx(blob)

    # CSV / TSV / plain text. Sniff the delimiter rather than assuming comma —
    # Indian exports are frequently semicolon-delimited.
    text = blob.decode("utf-8-sig", errors="replace")
    sample = text[:4096]
    try:
        dialect = csv.Sniffer().sniff(sample, delimiters=",;\t|")
    except csv.Error:
        dialect = csv.excel
    return [row for row in csv.reader(io.StringIO(text), dialect)]


def _read_docx(blob: bytes) -> list[list[Any]]:
    """DOCX: prefer real tables; fall back to one contact per paragraph.

    People genuinely do send a Word file with a list of numbers, one per line.
    Rejecting that would just create a support ticket.
    """
    try:
        import docx  # python-docx
    except ImportError:  # pragma: no cover
        return _read_docx_raw(blob)

    document = docx.Document(io.BytesIO(blob))
    rows: list[list[Any]] = []
    for table in document.tables:
        for row in table.rows:
            rows.append([cell.text for cell in row.cells])
    if rows:
        return rows

    rows.append(["name", "phone"])
    for paragraph in document.paragraphs:
        line = paragraph.text.strip()
        if not line:
            continue
        found = re.search(r"(\+?\d[\d\s\-()]{7,})", line)
        if found:
            number = found.group(1)
            rows.append([line.replace(number, "").strip(" -,:\t") or None, number])
    return rows


def _read_docx_raw(blob: bytes) -> list[list[Any]]:
    """No python-docx installed: pull text straight out of the OOXML zip."""
    import zipfile
    from xml.etree import ElementTree

    rows: list[list[Any]] = [["name", "phone"]]
    try:
        with zipfile.ZipFile(io.BytesIO(blob)) as archive:
            xml = archive.read("word/document.xml").decode("utf-8", errors="replace")
        text = re.sub(r"<[^>]+>", " ", xml)
    except (zipfile.BadZipFile, KeyError, ElementTree.ParseError):
        text = blob.decode("utf-8", errors="replace")
    for match in re.finditer(r"(\+?\d[\d\s\-()]{7,})", text):
        rows.append([None, match.group(1)])
    return rows


# =========================================================================== #
# main entry point
# =========================================================================== #
def parse_contacts(
    filename: str,
    content_type: str,
    blob: bytes,
    *,
    column_map: dict[str, str] | None = None,
    region: str = DEFAULT_REGION,
    dedupe: bool = True,
    max_rows: int | None = None,
) -> ParseResult:
    """Parse an uploaded contact list into validated, normalised rows."""
    limit = max_rows or settings.campaign_max_recipients
    result = ParseResult()

    try:
        matrix = _read_matrix(filename, content_type, blob)
    except Exception as exc:  # noqa: BLE001
        result.warnings.append(f"Could not read the file: {exc}")
        return result

    matrix = [row for row in matrix if any(str(c or "").strip() for c in row)]
    if not matrix:
        result.warnings.append("The file appears to be empty.")
        return result

    header_index = _find_header_row(matrix)
    headers = [str(c or "").strip() for c in matrix[header_index]]
    # Blank header cells still need a stable key so their data is not lost.
    headers = [h or f"column_{i + 1}" for i, h in enumerate(headers)]
    result.headers = headers

    mapping = guess_column_map(headers)
    if column_map:
        mapping.update({k: v for k, v in column_map.items() if v in headers})
    result.column_map = mapping

    if "phone" not in mapping and "whatsapp" not in mapping and "email" not in mapping:
        result.warnings.append(
            "No phone, WhatsApp or email column was detected. "
            "Map the columns manually before sending."
        )

    index_of = {header: i for i, header in enumerate(headers)}
    seen: set[str] = set()

    for offset, raw_row in enumerate(matrix[header_index + 1:], start=1):
        if len(result.rows) >= limit:
            result.warnings.append(
                f"Only the first {limit:,} rows were imported (limit reached)."
            )
            break

        def cell(canonical: str) -> Any:
            header = mapping.get(canonical)
            if header is None:
                return None
            position = index_of.get(header, -1)
            return raw_row[position] if 0 <= position < len(raw_row) else None

        parsed = ParsedRow(row_number=offset)
        parsed.name = (str(cell("name")).strip() if cell("name") else None) or None
        parsed.email = _valid_email(cell("email"))

        raw_phone = cell("phone") or cell("whatsapp")
        phone, error = normalise_phone(raw_phone, region)
        parsed.phone = phone
        parsed.whatsapp = phone
        if phone:
            parsed.country_code = phone[:3]

        # Extra columns survive for template placeholders.
        used = set(mapping.values())
        parsed.fields = {
            header: (str(raw_row[i]).strip() if i < len(raw_row) and raw_row[i] is not None else None)
            for header, i in index_of.items()
            if header not in used
        }
        parsed.fields = {k: v for k, v in parsed.fields.items() if v}

        if not phone and not parsed.email:
            parsed.is_valid = False
            parsed.invalid_reason = f"No contactable address ({error or 'no email'})"
        elif dedupe and phone:
            fingerprint = phone_fingerprint(phone)
            if fingerprint in seen:
                parsed.is_valid = False
                parsed.invalid_reason = "Duplicate of an earlier row"
                result.duplicates += 1
            else:
                seen.add(fingerprint)

        result.rows.append(parsed)

    result.total = len(result.rows)
    result.valid = sum(1 for r in result.rows if r.is_valid)
    result.invalid = result.total - result.valid - 0
    return result


def to_list_items(
    result: ParseResult, client_id: str, list_id: str, created_by: str
) -> list:
    """Materialise ParsedRows as LeadContactListItem ORM objects (encrypted)."""
    from ..models import LeadContactListItem

    items = []
    for row in result.rows:
        items.append(
            LeadContactListItem(
                ClientId=client_id,
                ListId=list_id,
                RowNumber=row.row_number,
                Name=row.name,
                PhoneEnc=encrypt_pii(row.phone),
                EmailEnc=encrypt_pii(row.email),
                WhatsAppEnc=encrypt_pii(row.whatsapp),
                PhoneHash=phone_fingerprint(row.phone),
                PhoneMasked=mask_phone(row.phone),
                CountryCode=row.country_code,
                IsValid=row.is_valid,
                InvalidReason=row.invalid_reason,
                FieldsJson=row.fields or None,
                CreatedBy=created_by,
            )
        )
    return items


def render_template(body: str, context: dict[str, Any]) -> str:
    """Substitute {{placeholders}} in a campaign message.

    Deliberately NOT Jinja: campaign bodies are authored by end users, and a
    real template engine there is a sandbox-escape surface. This is a literal
    replace of `{{key}}` with a string, nothing else — no logic, no attribute
    access, no imports. An unknown placeholder is emptied rather than left
    visible, so a customer never receives a raw `{{first_name}}`.
    """
    if not body:
        return ""

    def replace(match: re.Match) -> str:
        key = match.group(1).strip().lower()
        for candidate_key, value in context.items():
            if str(candidate_key).strip().lower() == key and value is not None:
                return str(value)
        return ""

    return re.sub(r"\{\{\s*([\w .-]+?)\s*\}\}", replace, body)
