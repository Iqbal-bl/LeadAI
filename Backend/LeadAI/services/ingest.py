"""
Turn an uploaded file into clean, retrievable chunks.

Chunking strategy (and why): paragraph-aware, ~900 characters, 150-character
overlap.

  * Paragraph-aware, not fixed-window, because company knowledge is written in
    sections — splitting mid-list turns "documents required: A, B, C" into two
    chunks that each answer the question wrongly.
  * ~900 chars ≈ 200 tokens: small enough that a retrieved chunk is mostly
    signal (a 4000-char chunk dilutes cosine similarity badly), large enough to
    hold a complete answer plus its heading.
  * 150-char overlap so an answer that straddles a boundary survives in at least
    one chunk. Overlap is trimmed to a word boundary so a chunk never begins
    mid-word, which would corrupt both the embedding and any excerpt shown as a
    citation.

The `_unwrap` step matters more than it looks: PDFs hard-wrap mid-sentence, and
without rejoining those lines every sentence-level score is computed on a
fragment.
"""
from __future__ import annotations

import io
import logging
import re

from ..config import settings

logger = logging.getLogger(__name__)

SUPPORTED = (".pdf", ".docx", ".txt", ".md", ".csv", ".json", ".xml", ".html")


def extract_text(filename: str, content_type: str, blob: bytes) -> str:
    name = (filename or "").lower()
    ctype = (content_type or "").lower()

    if name.endswith(".pdf") or "pdf" in ctype:
        return _from_pdf(blob)
    if name.endswith(".docx") or "wordprocessing" in ctype:
        return _from_docx(blob)
    if name.endswith((".html", ".htm")) or "html" in ctype:
        return _from_html(blob)
    return blob.decode("utf-8", errors="ignore")


def _from_pdf(blob: bytes) -> str:
    try:
        from pypdf import PdfReader
    except ImportError as exc:  # pragma: no cover
        raise RuntimeError("pypdf is required to ingest PDF files") from exc
    reader = PdfReader(io.BytesIO(blob))
    pages = []
    for page in reader.pages:
        try:
            pages.append(page.extract_text() or "")
        except Exception as exc:  # noqa: BLE001
            logger.warning("[LeadAI ingest] page extract failed: %s", exc)
    return "\n\n".join(pages)


def _from_docx(blob: bytes) -> str:
    try:
        import docx
    except ImportError as exc:  # pragma: no cover
        raise RuntimeError("python-docx is required to ingest .docx files") from exc
    document = docx.Document(io.BytesIO(blob))
    parts = [p.text for p in document.paragraphs if p.text.strip()]
    # Tables carry a lot of the value in rate cards / fee schedules, so flatten
    # them into pipe-separated rows rather than dropping them.
    for table in document.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts)


def _from_html(blob: bytes) -> str:
    text = blob.decode("utf-8", errors="ignore")
    text = re.sub(r"(?is)<(script|style).*?</\1>", " ", text)
    text = re.sub(r"(?s)<[^>]+>", "\n", text)
    return re.sub(r"&nbsp;?", " ", text)


def clean(text: str) -> str:
    text = (text or "").replace("\r", "\n")
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return _unwrap(text).strip()


def _unwrap(text: str) -> str:
    """Rejoin hard-wrapped lines while keeping real headings on their own line.

    Heuristic: a line is a continuation if the previous line is long (>45 chars)
    and does not end in terminal punctuation. A short line with no full stop is a
    heading, and keeping its break stops it being glued to the front of the first
    sentence of the section it introduces.
    """
    out: list[str] = []
    for line in text.split("\n"):
        stripped = line.strip()
        if not stripped:
            out.append("")
            continue
        prev = out[-1] if out else ""
        if prev and len(prev) > 45 and not prev.endswith((".", "!", "?", ":", ";")):
            out[-1] = f"{prev} {stripped}"
        else:
            out.append(stripped)
    return "\n".join(out)


def chunk(
    text: str,
    max_chars: int | None = None,
    overlap: int | None = None,
) -> list[str]:
    """Paragraph-aware splitter with word-boundary overlap."""
    max_chars = max_chars or settings.chunk_max_chars
    overlap = settings.chunk_overlap if overlap is None else overlap

    text = clean(text)
    if not text:
        return []

    def tail(buf: str) -> str:
        if not overlap:
            return ""
        window = buf[-overlap:]
        space = window.find(" ")
        return window[space + 1 :] if space != -1 else window

    paragraphs = [p.strip() for p in text.split("\n\n") if p.strip()]
    chunks: list[str] = []
    buffer = ""

    for para in paragraphs:
        if len(para) > max_chars:
            # A single oversized paragraph: fall back to sentence packing.
            for sentence in re.split(r"(?<=[.!?])\s+", para):
                if len(buffer) + len(sentence) + 1 > max_chars and buffer:
                    chunks.append(buffer.strip())
                    buffer = tail(buffer)
                buffer += " " + sentence
            continue
        if len(buffer) + len(para) + 2 > max_chars and buffer:
            chunks.append(buffer.strip())
            buffer = tail(buffer)
        buffer += "\n\n" + para

    if buffer.strip():
        chunks.append(buffer.strip())
    # Sub-20-char fragments are noise that pollutes retrieval with high-cosine,
    # zero-information hits.
    return [c for c in chunks if len(c) > 20]


def approx_tokens(text: str) -> int:
    """~4 chars per token. Good enough for budgeting a context window."""
    return max(1, len(text) // 4)
